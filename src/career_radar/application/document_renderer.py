"""把审核后的结构化正文确定性渲染为 DOCX，并可选转换为 PDF。"""

from __future__ import annotations

import hashlib
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from ..models import JobPosting
from .models import (
    ApplicationConfig,
    ApplicationDraftBundle,
    ApplicationProfile,
    ApplicationRun,
)

_SAFE_APPLICATION_ID = re.compile(r"^app-[A-Za-z0-9-]{8,100}$")
_ACCENT = "1F4E78"
_MUTED = "555555"
_BODY = "222222"
_CJK_FONT = "Microsoft YaHei"
_LATIN_FONT = "Arial"


class DocumentRenderError(RuntimeError):
    """DOCX 生成或强制 PDF 转换失败。"""


@dataclass(frozen=True)
class RenderedApplicationDocuments:
    output_dir: Path
    resume_docx: Path
    resume_pdf: Path | None
    cover_letter_docx: Path | None
    cover_letter_pdf: Path | None
    warnings: list[str] = field(default_factory=list)


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _set_run_font(
    run: Any,
    *,
    size: float,
    bold: bool = False,
    color: str = _BODY,
    italic: bool = False,
) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = _LATIN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), _LATIN_FONT)
    fonts.set(qn("w:hAnsi"), _LATIN_FONT)
    fonts.set(qn("w:eastAsia"), _CJK_FONT)


def _set_style_font(style: Any, *, size: float, bold: bool, color: str) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    style.font.name = _LATIN_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), _LATIN_FONT)
    fonts.set(qn("w:hAnsi"), _LATIN_FONT)
    fonts.set(qn("w:eastAsia"), _CJK_FONT)


def _configure_document(document: Any, *, compact: bool) -> None:
    """应用 compact_reference_guide 与中国校招 A4 命名覆盖。"""

    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Mm, Pt

    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    if compact:
        section.top_margin = Mm(13.5)
        section.bottom_margin = Mm(13.5)
        section.left_margin = Mm(16)
        section.right_margin = Mm(16)
        section.header_distance = Mm(8)
        section.footer_distance = Mm(8)
        body_size = 9.2
        body_after = 1.0
        line_spacing = 1.05
    else:
        section.top_margin = Mm(25.4)
        section.bottom_margin = Mm(25.4)
        section.left_margin = Mm(25.4)
        section.right_margin = Mm(25.4)
        section.header_distance = Mm(12.5)
        section.footer_distance = Mm(12.5)
        body_size = 11
        body_after = 6
        line_spacing = 1.5

    normal = document.styles["Normal"]
    _set_style_font(normal, size=body_size, bold=False, color=_BODY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(body_after)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = line_spacing

    heading = document.styles["Heading 1"]
    _set_style_font(heading, size=11.5 if compact else 13, bold=True, color=_ACCENT)
    heading.paragraph_format.space_before = Pt(4 if compact else 10)
    heading.paragraph_format.space_after = Pt(2 if compact else 5)
    heading.paragraph_format.keep_with_next = True

    if "Resume Item" not in document.styles:
        item_style = document.styles.add_style("Resume Item", WD_STYLE_TYPE.PARAGRAPH)
    else:
        item_style = document.styles["Resume Item"]
    _set_style_font(item_style, size=9.2 if compact else 11, bold=False, color=_BODY)
    item_style.paragraph_format.space_before = Pt(1.5 if compact else 3)
    item_style.paragraph_format.space_after = Pt(0.5 if compact else 2)
    item_style.paragraph_format.keep_with_next = True


def _add_bullet_numbering(document: Any) -> int:
    """创建真实 Word 编号定义，避免用 Unicode 字符伪造项目符号。"""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId"), "0"))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(element.get(qn("w:numId"), "0"))
        for element in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (("w:start", "1"), ("w:numFmt", "bullet"), ("w:lvlText", "•")):
        element = OxmlElement(tag)
        element.set(qn("w:val"), value)
        level.append(element)
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")
    level.append(level_justification)
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "270")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    paragraph_properties.extend([tabs, indent])
    level.append(paragraph_properties)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return num_id


def _add_bullet(document: Any, text: str, num_id: int, *, compact: bool) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1.2 if compact else 4)
    paragraph.paragraph_format.line_spacing = 1.05 if compact else 1.25
    properties = paragraph._p.get_or_add_pPr()
    number_properties = properties.get_or_add_numPr()
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    number_properties.extend([level, number])
    run = paragraph.add_run(text.strip())
    _set_run_font(run, size=9.2 if compact else 11)


def _add_section_heading(document: Any, title: str) -> None:
    document.add_paragraph(title, style="Heading 1")


def _add_item_header(
    document: Any,
    left: str,
    right: str | None,
    *,
    compact: bool,
) -> None:
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Mm

    paragraph = document.add_paragraph(style="Resume Item")
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Mm(178 if compact else 159.2), WD_TAB_ALIGNMENT.RIGHT
    )
    left_run = paragraph.add_run(left.strip())
    _set_run_font(left_run, size=9.4 if compact else 11, bold=True)
    if right:
        paragraph.add_run("\t")
        right_run = paragraph.add_run(right.strip())
        _set_run_font(right_run, size=8.8 if compact else 10, color=_MUTED)


def _plain_paragraph(
    document: Any,
    text: str,
    *,
    compact: bool,
    color: str = _BODY,
    bold: bool = False,
    italic: bool = False,
) -> Any:
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1.2 if compact else 6)
    paragraph.paragraph_format.line_spacing = 1.05 if compact else 1.5
    run = paragraph.add_run(text.strip())
    _set_run_font(
        run,
        size=9.2 if compact else 11,
        color=color,
        bold=bold,
        italic=italic,
    )
    return paragraph


def _scrub_docx_metadata(source: Path, destination: Path) -> None:
    """清除作者、自定义属性和 rsid，避免把本机信息带入投递文件。"""

    destination.parent.mkdir(parents=True, exist_ok=True)
    temporary = destination.with_name(f".{destination.name}.scrubbed")
    namespaces = {
        "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
        "dc": "http://purl.org/dc/elements/1.1/",
        "ct": "http://schemas.openxmlformats.org/package/2006/content-types",
        "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
    }
    with zipfile.ZipFile(source, "r") as archive_in, zipfile.ZipFile(
        temporary, "w", zipfile.ZIP_DEFLATED
    ) as archive_out:
        for member in archive_in.infolist():
            name = member.filename
            if name == "docProps/custom.xml":
                continue
            data = archive_in.read(member)
            if name == "docProps/core.xml":
                root = ElementTree.fromstring(data)
                for query in ("dc:creator", "cp:lastModifiedBy"):
                    element = root.find(query, namespaces)
                    if element is not None:
                        element.text = ""
                data = ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)
            elif name == "[Content_Types].xml":
                root = ElementTree.fromstring(data)
                for element in list(root):
                    if element.get("PartName") == "/docProps/custom.xml":
                        root.remove(element)
                data = ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)
            elif name == "_rels/.rels":
                root = ElementTree.fromstring(data)
                for element in list(root):
                    if element.get("Target") == "docProps/custom.xml":
                        root.remove(element)
                data = ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)
            elif name.startswith("word/") and name.endswith(".xml"):
                root = ElementTree.fromstring(data)
                for element in root.iter():
                    for attribute in list(element.attrib):
                        if attribute.rsplit("}", 1)[-1].startswith("rsid"):
                            del element.attrib[attribute]
                data = ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)
            archive_out.writestr(member, data)
    os.replace(temporary, destination)


def _save_document(document: Any, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    building = target.with_name(f".{target.name}.building.docx")
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.save(building)
    try:
        _scrub_docx_metadata(building, target)
    finally:
        building.unlink(missing_ok=True)


class ApplicationDocumentRenderer:
    """只消费冻结快照和终稿；不会再调用 LLM，也不会修改终稿事实。"""

    def __init__(self, config: ApplicationConfig) -> None:
        self.config = config

    def _output_directory(self, application_id: str) -> Path:
        if not _SAFE_APPLICATION_ID.fullmatch(application_id):
            raise DocumentRenderError("申请任务 ID 格式异常，拒绝创建输出目录")
        root = self.config.output_dir.expanduser().resolve()
        target = (root / application_id).resolve()
        if target.parent != root:
            raise DocumentRenderError("申请材料输出路径越界")
        target.mkdir(parents=True, exist_ok=True)
        return target

    @staticmethod
    def _new_document() -> Any:
        try:
            from docx import Document
        except ImportError as exc:
            raise DocumentRenderError(
                "未安装 python-docx，请执行 python -m pip install -e ."
            ) from exc
        return Document()

    def _build_resume(
        self,
        profile: ApplicationProfile,
        drafts: ApplicationDraftBundle,
    ) -> Any:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        document = self._new_document()
        _configure_document(document, compact=True)
        bullet_id = _add_bullet_numbering(document)
        resume = drafts.resume

        name = document.add_paragraph()
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name.paragraph_format.space_after = Pt(1)
        _set_run_font(name.add_run(profile.contact.name), size=19, bold=True, color=_ACCENT)

        headline = document.add_paragraph()
        headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        headline.paragraph_format.space_after = Pt(1.5)
        _set_run_font(headline.add_run(resume.headline), size=10.2, bold=True, color=_MUTED)

        contact_values = [
            profile.contact.phone,
            profile.contact.email,
            profile.contact.location,
            profile.contact.github,
            profile.contact.linkedin,
        ]
        contact = document.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.paragraph_format.space_after = Pt(3)
        _set_run_font(
            contact.add_run(" | ".join(value for value in contact_values if value)),
            size=8.2,
            color=_MUTED,
        )

        _add_section_heading(document, "个人简介")
        _plain_paragraph(document, resume.professional_summary, compact=True)

        if resume.skills:
            _add_section_heading(document, "专业技能")
            _plain_paragraph(
                document,
                " · ".join(item.text for item in resume.skills),
                compact=True,
            )

        if resume.education:
            _add_section_heading(document, "教育经历")
            for item in resume.education:
                _add_item_header(document, item.institution, item.period, compact=True)
                _plain_paragraph(
                    document,
                    " | ".join(value for value in (item.degree, item.major) if value),
                    compact=True,
                    color=_MUTED,
                )
                for highlight in item.highlights:
                    _add_bullet(document, highlight, bullet_id, compact=True)

        if resume.experiences:
            _add_section_heading(document, "实习与工作经历")
            for item in resume.experiences:
                label = " | ".join(value for value in (item.organization, item.role) if value)
                _add_item_header(document, label, item.period, compact=True)
                if item.location:
                    _plain_paragraph(document, item.location, compact=True, color=_MUTED)
                for bullet in item.bullets:
                    _add_bullet(document, bullet, bullet_id, compact=True)

        if resume.projects:
            _add_section_heading(document, "项目经历")
            for item in resume.projects:
                _add_item_header(document, item.name, item.period, compact=True)
                if item.summary:
                    _plain_paragraph(document, item.summary, compact=True, color=_MUTED)
                for bullet in item.bullets:
                    _add_bullet(document, bullet, bullet_id, compact=True)
                if item.technologies:
                    _plain_paragraph(
                        document,
                        f"技术栈：{' / '.join(item.technologies)}",
                        compact=True,
                        color=_MUTED,
                    )

        if resume.awards:
            _add_section_heading(document, "奖项与荣誉")
            for item in resume.awards:
                _add_bullet(document, item.text, bullet_id, compact=True)

        if resume.leadership:
            _add_section_heading(document, "校园与组织经历")
            for item in resume.leadership:
                _add_bullet(document, item.text, bullet_id, compact=True)
        return document

    def _build_cover_letter(
        self,
        job: JobPosting,
        profile: ApplicationProfile,
        drafts: ApplicationDraftBundle,
        generated_date: str,
    ) -> Any:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        cover = drafts.cover_letter
        if cover is None:
            raise DocumentRenderError("没有求职信终稿")
        document = self._new_document()
        _configure_document(document, compact=False)

        identity = document.add_paragraph()
        identity.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        identity.paragraph_format.space_after = Pt(2)
        _set_run_font(identity.add_run(profile.contact.name), size=13, bold=True, color=_ACCENT)
        contact = document.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        contact.paragraph_format.space_after = Pt(16)
        _set_run_font(
            contact.add_run(f"{profile.contact.phone} | {profile.contact.email}"),
            size=9,
            color=_MUTED,
        )

        date = document.add_paragraph()
        date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date.paragraph_format.space_after = Pt(16)
        _set_run_font(date.add_run(generated_date), size=10, color=_MUTED)

        subject = document.add_paragraph()
        subject.paragraph_format.space_after = Pt(16)
        _set_run_font(subject.add_run(cover.subject), size=15, bold=True, color=_ACCENT)
        _plain_paragraph(document, f"申请岗位：{job.company} · {job.title}", compact=False, color=_MUTED)
        _plain_paragraph(document, cover.salutation, compact=False)
        for value in cover.paragraphs:
            paragraph = _plain_paragraph(document, value, compact=False)
            paragraph.paragraph_format.first_line_indent = Pt(22)
        _plain_paragraph(document, cover.closing, compact=False)
        signature = document.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_run_font(signature.add_run(profile.contact.name), size=11, bold=True)
        return document

    def _locate_office(self) -> str | None:
        configured = Path(self.config.libreoffice_command).expanduser()
        if configured.is_file():
            return str(configured.resolve())
        discovered = shutil.which(self.config.libreoffice_command)
        if discovered:
            return discovered
        for candidate in (
            Path("C:/Program Files/LibreOffice/program/soffice.exe"),
            Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),
        ):
            if candidate.is_file():
                return str(candidate)
        return None

    def _convert_pdf(self, docx_path: Path, output_pdf: Path) -> tuple[Path | None, str | None]:
        if self.config.pdf_mode == "never":
            output_pdf.unlink(missing_ok=True)
            return None, None
        office = self._locate_office()
        if office is None:
            output_pdf.unlink(missing_ok=True)
            message = "未检测到 LibreOffice，已保留 DOCX，未自动生成 PDF"
            if self.config.pdf_mode == "always":
                raise DocumentRenderError(message)
            return None, message
        output_pdf.unlink(missing_ok=True)
        with tempfile.TemporaryDirectory(prefix="career-radar-lo-", dir=docx_path.parent) as raw:
            temporary = Path(raw)
            profile = temporary / "profile"
            profile.mkdir()
            command = [
                office,
                "--headless",
                f"-env:UserInstallation={profile.as_uri()}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(temporary),
                str(docx_path),
            ]
            creation_flags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                timeout=180,
                check=False,
                creationflags=creation_flags,
            )
            converted = temporary / f"{docx_path.stem}.pdf"
            if result.returncode != 0 or not converted.is_file() or converted.stat().st_size == 0:
                detail = (result.stderr or result.stdout or "未知转换错误").strip()[-1000:]
                raise DocumentRenderError(f"LibreOffice PDF 转换失败：{detail}")
            building = output_pdf.with_name(f".{output_pdf.name}.building")
            shutil.copy2(converted, building)
            os.replace(building, output_pdf)
        return output_pdf, None

    def render(
        self,
        run: ApplicationRun,
        job: JobPosting,
        profile: ApplicationProfile,
        drafts: ApplicationDraftBundle,
        generated_date: str,
    ) -> RenderedApplicationDocuments:
        output_dir = self._output_directory(run.id)
        warnings: list[str] = []

        resume_docx = output_dir / "resume.docx"
        _save_document(self._build_resume(profile, drafts), resume_docx)
        resume_pdf, warning = self._convert_pdf(resume_docx, output_dir / "resume.pdf")
        if warning:
            warnings.append(warning)

        cover_docx = None
        cover_pdf = None
        if drafts.cover_letter is not None:
            cover_docx = output_dir / "cover_letter.docx"
            _save_document(
                self._build_cover_letter(job, profile, drafts, generated_date),
                cover_docx,
            )
            cover_pdf, warning = self._convert_pdf(
                cover_docx, output_dir / "cover_letter.pdf"
            )
            if warning and warning not in warnings:
                warnings.append(warning)

        return RenderedApplicationDocuments(
            output_dir=output_dir,
            resume_docx=resume_docx,
            resume_pdf=resume_pdf,
            cover_letter_docx=cover_docx,
            cover_letter_pdf=cover_pdf,
            warnings=warnings,
        )

    def existing(
        self,
        run: ApplicationRun,
        drafts: ApplicationDraftBundle,
    ) -> RenderedApplicationDocuments:
        """从确定性文件名恢复 verifying 断点，不信任数据库中的任意路径。"""

        output_dir = self._output_directory(run.id)
        resume_docx = output_dir / "resume.docx"
        if not resume_docx.is_file():
            raise DocumentRenderError("恢复文档验证时找不到 resume.docx")
        resume_pdf_path = output_dir / "resume.pdf"
        cover_docx_path = output_dir / "cover_letter.docx"
        cover_pdf_path = output_dir / "cover_letter.pdf"
        cover_docx = cover_docx_path if drafts.cover_letter is not None else None
        if cover_docx is not None and not cover_docx.is_file():
            raise DocumentRenderError("恢复文档验证时找不到 cover_letter.docx")
        return RenderedApplicationDocuments(
            output_dir=output_dir,
            resume_docx=resume_docx,
            resume_pdf=resume_pdf_path if resume_pdf_path.is_file() else None,
            cover_letter_docx=cover_docx,
            cover_letter_pdf=(
                cover_pdf_path
                if drafts.cover_letter is not None and cover_pdf_path.is_file()
                else None
            ),
        )
