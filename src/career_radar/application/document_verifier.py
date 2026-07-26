"""对申请 DOCX/PDF 做确定性的结构、完整性、元数据和页数校验。"""

from __future__ import annotations

import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from ..models import JobPosting
from .document_renderer import RenderedApplicationDocuments, file_sha256
from .models import (
    ApplicationConfig,
    ApplicationDraftBundle,
    ApplicationProfile,
    ApplicationRun,
    ApplicationVerification,
    DocumentVerification,
    VerificationIssue,
)


def _normalized(value: str) -> str:
    return re.sub(r"\s+", "", value)


def _resume_expected_texts(
    profile: ApplicationProfile,
    drafts: ApplicationDraftBundle,
) -> list[str]:
    resume = drafts.resume
    values = [
        profile.contact.name,
        profile.contact.phone,
        profile.contact.email,
        profile.contact.location,
        profile.contact.github,
        profile.contact.linkedin,
        resume.headline,
        resume.professional_summary,
    ]
    values.extend(item.text for item in resume.skills)
    for item in resume.education:
        values.extend([item.institution, item.degree, item.major, item.period, *item.highlights])
    for item in resume.experiences:
        values.extend(
            [item.organization, item.role, item.period, item.location, *item.bullets]
        )
    for item in resume.projects:
        values.extend(
            [item.name, item.period, item.summary, *item.bullets, *item.technologies]
        )
    values.extend(item.text for item in resume.awards)
    values.extend(item.text for item in resume.leadership)
    return list(dict.fromkeys(value.strip() for value in values if value and value.strip()))


def _cover_expected_texts(
    job: JobPosting,
    profile: ApplicationProfile,
    drafts: ApplicationDraftBundle,
) -> list[str]:
    cover = drafts.cover_letter
    if cover is None:
        return []
    values = [
        profile.contact.name,
        profile.contact.phone,
        profile.contact.email,
        job.company,
        job.title,
        cover.subject,
        cover.salutation,
        *cover.paragraphs,
        cover.closing,
    ]
    return list(dict.fromkeys(value.strip() for value in values if value.strip()))


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("未安装 python-docx，无法校验生成的 Word 文档") from exc
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    for section in document.sections:
        values.extend(paragraph.text for paragraph in section.header.paragraphs)
        values.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(values)


def _metadata_issues(path: Path, document_kind: str) -> list[VerificationIssue]:
    issues: list[VerificationIssue] = []
    with zipfile.ZipFile(path, "r") as archive:
        corrupt_member = archive.testzip()
        if corrupt_member:
            issues.append(
                VerificationIssue(
                    severity="error",
                    code="docx_zip_corrupt",
                    document_kind=document_kind,
                    message="DOCX 压缩包存在损坏成员",
                )
            )
        if "docProps/custom.xml" in archive.namelist():
            issues.append(
                VerificationIssue(
                    severity="error",
                    code="custom_metadata_present",
                    document_kind=document_kind,
                    message="DOCX 仍包含自定义元数据",
                )
            )
        if "docProps/core.xml" in archive.namelist():
            root = ElementTree.fromstring(archive.read("docProps/core.xml"))
            for element in root.iter():
                if element.tag.rsplit("}", 1)[-1] in {"creator", "lastModifiedBy"} and (
                    element.text or ""
                ).strip():
                    issues.append(
                        VerificationIssue(
                            severity="error",
                            code="author_metadata_present",
                            document_kind=document_kind,
                            message="DOCX 作者元数据未清理",
                        )
                    )
                    break
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            root = ElementTree.fromstring(archive.read(name))
            if any(
                attribute.rsplit("}", 1)[-1].startswith("rsid")
                for element in root.iter()
                for attribute in element.attrib
            ):
                issues.append(
                    VerificationIssue(
                        severity="error",
                        code="revision_metadata_present",
                        document_kind=document_kind,
                        message="DOCX 仍包含 Word 修订会话标识",
                    )
                )
                break
    return issues


def _layout_issues(
    path: Path,
    document_kind: str,
    expected_bullet_count: int,
) -> list[VerificationIssue]:
    """审计不依赖 Word/LibreOffice 的关键 OOXML 版式契约。"""

    issues: list[VerificationIssue] = []
    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": word_namespace}
    attribute = lambda name: f"{{{word_namespace}}}{name}"  # noqa: E731
    with zipfile.ZipFile(path, "r") as archive:
        document = ElementTree.fromstring(archive.read("word/document.xml"))
        page = document.find(".//w:sectPr/w:pgSz", ns)
        margins = document.find(".//w:sectPr/w:pgMar", ns)
        page_width = int(page.get(attribute("w"), "0")) if page is not None else 0
        page_height = int(page.get(attribute("h"), "0")) if page is not None else 0
        if abs(page_width - 11906) > 5 or abs(page_height - 16838) > 5:
            issues.append(
                VerificationIssue(
                    severity="error",
                    code="page_size_not_a4",
                    document_kind=document_kind,
                    message="文档页面不是预期的 A4 纵向尺寸",
                )
            )
        expected_margin = 907 if document_kind == "resume" else 1440
        if margins is None or any(
            abs(int(margins.get(attribute(side), "0")) - expected_margin) > 6
            for side in ("left", "right")
        ):
            issues.append(
                VerificationIssue(
                    severity="error",
                    code="horizontal_margins_invalid",
                    document_kind=document_kind,
                    message="文档左右页边距不符合选定样式契约",
                )
            )

        numbered_paragraphs = document.findall(".//w:p/w:pPr/w:numPr", ns)
        if len(numbered_paragraphs) < expected_bullet_count:
            issues.append(
                VerificationIssue(
                    severity="error",
                    code="real_bullets_missing",
                    document_kind=document_kind,
                    message="部分列表未使用真实 Word 编号定义",
                )
            )
        for paragraph in document.findall(".//w:p", ns):
            text = "".join(element.text or "" for element in paragraph.findall(".//w:t", ns))
            if text.lstrip().startswith(("•", "- ")):
                issues.append(
                    VerificationIssue(
                        severity="error",
                        code="fake_bullet_detected",
                        document_kind=document_kind,
                        message="检测到以字符伪造的项目符号",
                    )
                )
                break
        if document_kind == "resume" and document.find(".//w:tbl", ns) is not None:
            issues.append(
                VerificationIssue(
                    severity="error",
                    code="layout_table_detected",
                    document_kind=document_kind,
                    message="简历使用了布局表格，不符合当前 ATS 友好样式",
                )
            )
        styles = ElementTree.fromstring(archive.read("word/styles.xml"))
        heading = styles.find(".//w:style[@w:styleId='Heading1']", ns)
        if document_kind == "resume" and heading is None:
            issues.append(
                VerificationIssue(
                    severity="error",
                    code="heading_style_missing",
                    document_kind=document_kind,
                    message="简历章节没有使用真实 Heading 1 样式",
                )
            )
    return issues


class ApplicationDocumentVerifier:
    def __init__(self, config: ApplicationConfig) -> None:
        self.config = config

    def _verify_one(
        self,
        document_kind: str,
        docx_path: Path,
        pdf_path: Path | None,
        expected_texts: list[str],
        page_target: int,
        expected_bullet_count: int,
    ) -> DocumentVerification:
        if not docx_path.is_file() or docx_path.stat().st_size == 0:
            raise RuntimeError(f"缺少生成的 {document_kind} DOCX")
        issues = [
            *_metadata_issues(docx_path, document_kind),
            *_layout_issues(docx_path, document_kind, expected_bullet_count),
        ]
        actual_text = _normalized(_extract_docx_text(docx_path))
        missing = [value for value in expected_texts if _normalized(value) not in actual_text]
        if missing:
            issues.append(
                VerificationIssue(
                    severity="error",
                    code="expected_text_missing",
                    document_kind=document_kind,
                    message=f"生成文档缺少 {len(missing)} 项终稿或联系方式内容",
                )
            )

        pdf_hash = None
        pdf_bytes = None
        page_count = None
        if pdf_path is not None:
            if not pdf_path.is_file() or pdf_path.stat().st_size == 0:
                issues.append(
                    VerificationIssue(
                        severity="error",
                        code="pdf_missing",
                        document_kind=document_kind,
                        message="PDF 路径已登记但文件不存在或为空",
                    )
                )
            else:
                with pdf_path.open("rb") as handle:
                    if handle.read(5) != b"%PDF-":
                        issues.append(
                            VerificationIssue(
                                severity="error",
                                code="pdf_signature_invalid",
                                document_kind=document_kind,
                                message="PDF 文件头无效",
                            )
                        )
                try:
                    from pypdf import PdfReader

                    page_count = len(PdfReader(str(pdf_path)).pages)
                except Exception as exc:
                    issues.append(
                        VerificationIssue(
                            severity="error",
                            code="pdf_parse_failed",
                            document_kind=document_kind,
                            message=f"PDF 无法解析：{type(exc).__name__}",
                        )
                    )
                pdf_hash = file_sha256(pdf_path)
                pdf_bytes = pdf_path.stat().st_size
                if page_count is not None and page_count > page_target:
                    issues.append(
                        VerificationIssue(
                            severity="error",
                            code="page_target_exceeded",
                            document_kind=document_kind,
                            message=f"PDF 共 {page_count} 页，超过目标 {page_target} 页",
                        )
                    )
        elif self.config.pdf_mode == "always":
            issues.append(
                VerificationIssue(
                    severity="error",
                    code="pdf_required",
                    document_kind=document_kind,
                    message="配置要求必须生成 PDF，但当前没有 PDF",
                )
            )
        elif self.config.pdf_mode == "auto":
            issues.append(
                VerificationIssue(
                    severity="warning",
                    code="pdf_converter_unavailable",
                    document_kind=document_kind,
                    message="未生成 PDF，无法自动验证最终页数和视觉布局",
                )
            )

        return DocumentVerification(
            document_kind=document_kind,
            docx_sha256=file_sha256(docx_path),
            docx_bytes=docx_path.stat().st_size,
            pdf_sha256=pdf_hash,
            pdf_bytes=pdf_bytes,
            pdf_page_count=page_count,
            expected_text_checks=len(expected_texts),
            issues=issues,
        )

    def verify(
        self,
        run: ApplicationRun,
        job: JobPosting,
        profile: ApplicationProfile,
        drafts: ApplicationDraftBundle,
        rendered: RenderedApplicationDocuments,
        generated_at: str,
    ) -> ApplicationVerification:
        documents = [
            self._verify_one(
                "resume",
                rendered.resume_docx,
                rendered.resume_pdf,
                _resume_expected_texts(profile, drafts),
                run.resume_page_target,
                sum(len(item.highlights) for item in drafts.resume.education)
                + sum(len(item.bullets) for item in drafts.resume.experiences)
                + sum(len(item.bullets) for item in drafts.resume.projects)
                + len(drafts.resume.awards)
                + len(drafts.resume.leadership),
            )
        ]
        if drafts.cover_letter is not None:
            if rendered.cover_letter_docx is None:
                raise RuntimeError("求职信终稿存在，但没有生成求职信 DOCX")
            documents.append(
                self._verify_one(
                    "cover_letter",
                    rendered.cover_letter_docx,
                    rendered.cover_letter_pdf,
                    _cover_expected_texts(job, profile, drafts),
                    1,
                    0,
                )
            )
        passed = not any(
            issue.severity == "error"
            for document in documents
            for issue in document.issues
        )
        return ApplicationVerification(
            application_id=run.id,
            passed=passed,
            documents=documents,
            generated_at=generated_at,
        )
